import os
import time

import threadpool as tp
import re
import yaml
from docx import Document
import random
from tkinter.filedialog import askdirectory
from temp_storage import read_file_to_temp_list


def delInTable(tables, kwd='__删除整行__'):
    if tables == []:
        return
    for t, table in enumerate(tables):
        for r, row in enumerate(table.rows):  # 遍历表格的行
            for c, cell in enumerate(row.cells):  # 遍历每一行的列
                text = cell.text
                if kwd in text:
                    row._element.getparent().remove(row._element)
                    break


def del_rows(doc_path, kwd='__删除整行__'):
    if doc_path is None:
        return
    document = Document(doc_path)
    delInTable(document.tables, kwd)
    for t, table in enumerate(document.tables):
        for r, row in enumerate(table.rows):  # 遍历表格的行
            for c, cell in enumerate(row.cells):  # 遍历每一行的列
                delInTable(cell.tables, kwd)
    document.save(doc_path)


def delTextInTable(tables, kwd='__删除文本__'):
    if tables == []:
        return
    del_list = []
    for t, table in enumerate(tables):  # 遍历所有表格，这个模板里有3个表格
        for r, row in enumerate(table.rows):  # 遍历表格的行
            for c, cell in enumerate(row.cells):  # 遍历每一行的列
                text = cell.text
                if kwd in text:
                    del_list.append((t, r))
                    break

    for t, r in set(del_list):
        row = tables[t].rows[r]
        for cell in row.cells:
            cell.text = ""


def del_text(doc_path, kwd='__删除文本__'):
    if doc_path is None:
        return
    document = Document(doc_path)
    delTextInTable(document.tables, kwd)
    for t, table in enumerate(document.tables):  # 遍历所有表格，这个模板里有3个表格
        for r, row in enumerate(table.rows):  # 遍历表格的行
            for c, cell in enumerate(row.cells):  # 遍历每一行的列
                delTextInTable(cell.tables, kwd)

    document.save(doc_path)


def generate_dictionary(info_dict):
    deltext = get_dictionary("delText.yml")
    delline = get_dictionary('delLine.yml')
    white_list = get_dictionary('white_list.yml')
    dictionary = get_dictionary("tagList.yml")
    dictionary['template_id'] = info_dict['template_id']
    for key in dictionary.keys():
        if key in info_dict and info_dict[key] != '无' and info_dict[key] != '否':
            dictionary[key] = info_dict[key]
        elif key[2:-2] in info_dict and info_dict[key[2:-2]] != '无' and info_dict[key[2:-2]] != '否':
            dictionary[key] = info_dict[key[2:-2]]
        elif '随机日期' in key:
            continue
        else:
            dictionary[key] = '否'
        if key in ['__特殊过程焊接__', '__特殊过程搅拌__', '__特殊过程薄膜金属化__', '__特殊过程混合__',
                   '__特殊过程挤出__']:
            if key in info_dict and info_dict[key] == '是':
                dictionary[key] = ''

        if key in delline and dictionary[key] == delline[key]:
            dictionary[key] = '__删除整行__'
        if key in deltext and dictionary[key] == deltext[key]:
            dictionary[key] = '__删除文本__'
        if key in white_list:
            # 复评模式中的白名单词条，根据这种格式定位某些敏感词条的位置
            dictionary[key] = ' {} '.format(dictionary[key])
    if dictionary['__有无外包过程__'] == '否':
        dictionary['__外包过程表述__'] = ''
        dictionary['__有无外包过程__'] = '无'
    elif dictionary['__有无外包过程__'] == '是':
        dictionary['__有无外包过程__'] = '有'

    for key in dictionary.keys():
        if dictionary[key] == '否':
            dictionary[key] = ''
    for key in ['__1月随机日期__', '__2月随机日期__', '__3月随机日期__', '__4月随机日期__', '__5月随机日期__',
                '__6月随机日期__',
                '__7月随机日期__', '__8月随机日期__', '__9月随机日期__', '__10月随机日期__', '__11月随机日期__',
                '__12月随机日期__']:
        info_dict[key] = dictionary[key]
    return dictionary


def generate_old2new_dic(old2new: dict, info_dic: dict):
    olddict = read_file_to_temp_list(info_dic['tmp'])
    deltext = get_dictionary("delText.yml")
    delline = get_dictionary('delLine.yml')
    white_list = get_dictionary('white_list.yml')
    for k, v in info_dic.items():
        tmpk = k.replace('__', '')
        if tmpk in olddict and v != olddict[tmpk]:
            if k in white_list:
                old2new[' {} '.format(olddict[tmpk])] = v
                continue
            else:
                old2new[olddict[tmpk]] = v

            if k in deltext and v == '':
                old2new[olddict[tmpk]] = '__删除文本__'
                continue
            if k in delline and v == '':
                old2new[olddict[tmpk]] = '__删除整行__'
                continue


def get_dictionary(yaml_file):
    """
    从yamlFile中读取替换字典
    """
    with open(yaml_file, 'r', encoding='utf8') as y:
        return yaml.safe_load(y.read())


def replace_in_runs(paragraph, replace_dict: dict, paragraph_type):
    """
    将paragraph中所有runs根据replaceDict进行替换
    """
    text = ""
    if paragraph_type == "plain" or "header":
        for run in paragraph.runs:
            if text == "":
                if run.text.count("__") % 2 != 0:  # 发生了截断
                    text = run.text
                    run.text = ""
                    continue
            else:
                run.text = text + run.text
                if run.text.count("__") % 2 != 0:  # 发生了截断
                    text = run.text
                    run.text = ""
                    continue
            text = ""
            for old, new in replace_dict.items():
                # 进行替换
                if old == ' {} '.format(run.text):
                    # 监督复评模式下对于白名单的替换策略，初审的时候关键字已经通过两个空格进行隔断，
                    # 复评的时候先将键的格式设为" XXX "然后value保持原来的格式，由于空格已经将原有的关键词单独分成一个run所以可以对这个run进行直接替换
                    run.text = new
                    break
                if old in run.text:
                    if '版本号' in old:
                        print('WHZ---------检测到版本号关键词')
                        print(text)
                    if "随机日期" in old:
                        if '-' in new:
                            start, end = new.split('-')
                            new = str(random.randint(int(start), int(end)))
                            replace_dict[old] = new
                    if "__是否有空压机__" in old and new == '':
                        return 1
                    if "产品批号__" in old or "产品订货时间__" in old:
                        match_obj = re.search(r'__[\u4e00-\u9fa5|0-9]*__', old, re.M | re.I)
                        key1 = match_obj.group()[:12] + "批号__"
                        key2 = match_obj.group()[:12] + "订货时间__"
                        if replace_dict[key1] == '' and replace_dict[key2] == '':
                            return 1
                    run.text = run.text.replace(old, new)
    else:
        if paragraph_type == "table":
            for run in paragraph.runs:
                text += run.text.strip()
            for old, new in replace_dict.items():
                # 进行替换
                if old in text:
                    if '版本号' in old:
                        print('WHZ---------检测到版本号关键词')
                        print(text)
                    if "随机日期" in old:
                        if '-' in new:
                            start, end = new.split('-')
                            new = str(random.randint(int(start), int(end)))
                            replace_dict[old] = new
                    if "__是否有空压机__" in old and new == '':
                        return 1
                    if "产品批号__" in old or "产品订货时间__" in old:
                        match_obj = re.search(r'__[\u4e00-\u9fa5|0-9]*__', old, re.M | re.I)
                        key1 = match_obj.group()[:12] + "批号__"
                        key2 = match_obj.group()[:12] + "订货时间__"
                        if replace_dict[key1] == '' and replace_dict[key2] == '':
                            return 1
                    text = text.replace(old, new)
            if text != "":
                paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
    return 0


def replace_in_table(tables, replace_dict):
    if not tables:
        return 0
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    flag1 = replace_in_runs(paragraph, replace_dict, "table")
                flag2 = replace_in_table(cell.tables, replace_dict)
    return flag1 + flag2


def replace_in(docx_file, workdir, root_dir, res_home, replace_dict: dict):
    """
    将docxFile按replaceDict进行全局替换
    """
    if docx_file[-4:] == '.doc':
        print('无法加载doc文件： ', docx_file, '\n请转换成docx类型！')
        return
    else:
        print('正在处理：', docx_file)

    if docx_file.split('/')[-1][:2] == '.~':
        return

    document = Document(docx_file)
    res_path = os.path.join(workdir, res_home,
                            os.path.relpath(docx_file, root_dir))
    res_dir = os.path.dirname(res_path)
    if not os.path.exists(res_dir):
        os.makedirs(res_dir)

    # 页眉
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            flag = replace_in_runs(paragraph, replace_dict, "header")
            if flag:
                return None

    # 文本段落
    for paragraph in document.paragraphs:
        flag = replace_in_runs(paragraph, replace_dict, "plain")
        if flag:
            return None

    # 表格
    flag = replace_in_table(document.tables, replace_dict)
    if flag > 0:
        return None

    print(res_path)
    document.save(res_path)
    return res_path


def reset_date(dic, inf):
    for key in ['__1月随机日期__', '__2月随机日期__', '__3月随机日期__', '__4月随机日期__', '__5月随机日期__',
                '__6月随机日期__',
                '__7月随机日期__', '__8月随机日期__', '__9月随机日期__', '__10月随机日期__', '__11月随机日期__',
                '__12月随机日期__']:
        if key in dic:
            dic[key] = inf[key]


def replace_all(path: str, workdir: str, root_dir: str, res_home: str, replace_dict: dict, del_line_dict: dict,
                del_text_dict: dict, info_dict: dict):
    """
    将path路径下所有word文件按replaceDict进行全局替换（已更换）
    """
    directories = os.listdir(path)
    folders = [os.path.join(path, i) for i in directories if
               os.path.isdir(os.path.join(path, i)) and not i.startswith('.')]  # 所有非隐藏文件夹
    for folder in folders:
        replace_all(folder, workdir, root_dir, res_home, replace_dict, del_line_dict, del_text_dict, info_dict)
    files = [os.path.join(path, i) for i in directories if
             os.path.isfile(os.path.join(path, i)) and os.path.splitext(i)[-1] == '.docx']  # 所有word文件
    for file in files:
        respath = replace_in(file, workdir, root_dir, res_home, replace_dict)
        if respath is None:
            os.remove(file)
        del_rows(respath)
        del_text(respath)
        reset_date(replace_dict, info_dict)


def get_all_files(path: str):
    directories = os.listdir(path)
    folders = [os.path.join(path, i) for i in directories if
               os.path.isdir(os.path.join(path, i)) and not i.startswith('.')]
    files_list = []
    for folder in folders:
        files_list.extend(get_all_files(folder))
    files_list.extend([os.path.join(path, i) for i in directories if
                       os.path.isfile(os.path.join(path, i)) and os.path.splitext(i)[-1] == '.docx'])
    return files_list


def file_deal_process(file, workdir, root_dir, res_home, replace_dict, info_dict):
    respath = replace_in(file, workdir, root_dir, res_home, replace_dict)
    if respath is None:
        os.remove(file)
    del_rows(respath)
    del_text(respath)
    reset_date(replace_dict, info_dict)


def exam_in_runs(paragraph, paragraph_type):
    """
    将paragraph中所有runs根据replaceDict进行替换
    """
    text = ""
    if paragraph_type == "plain" or "header":
        for run in paragraph.runs:
            if run.text.count('__') > 0:
                return 1
    else:
        if paragraph_type == "table":
            for run in paragraph.runs:
                text += run.text.strip()
                if '__' in text:
                    return 1
    return 0


def exam_in(docx_file):
    """
    将docxFile进行是否有__检查
    """
    if docx_file[-4:] == '.doc':
        print('无法加载doc文件： ', docx_file, '\n请转换成docx类型！')
        return
    else:
        print('正在检查：', docx_file)

    if docx_file.split('/')[-1][:2] == '.~':
        return

    document = Document(docx_file)

    # 页眉
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            if exam_in_runs(paragraph, "header") == 1:
                return docx_file

    # 文本段落
    for paragraph in document.paragraphs:
        if exam_in_runs(paragraph, "plain") == 1:
            return docx_file

    # 表格
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if exam_in_runs(paragraph, "table") == 1:
                        return docx_file


def exam_all(path: str, depth: int):
    """
    将path路径下所有word文件进行是否有__检查
    """
    directories = os.listdir(path)
    undeal_files = []
    folders = [os.path.join(path, i) for i in directories if
               os.path.isdir(os.path.join(path, i)) and not i.startswith('.')]  # 所有非隐藏文件夹
    for folder in folders:
        undeal_files += exam_all(folder, 1)
    files = [os.path.join(path, i) for i in directories if
             os.path.isfile(os.path.join(path, i)) and os.path.splitext(i)[-1] in ['.doc', '.docx']]  # 所有word文件
    for file in files:
        undeal_files.append(exam_in(file))
    if depth == 0:
        with open(path + '/undeal.txt', 'w') as f:
            for line in undeal_files:
                if line:
                    f.write(line + '\n')
    return undeal_files


def replace_process(info_dict, re=False, page2=False, old2new=None):
    dictionary = generate_dictionary(info_dict)
    store_dir = askdirectory()
    year = dictionary['__记录年份__'] if '__记录年份__' in dictionary else '未知年份'
    res_home = year + dictionary['__企业名称__'] + dictionary['template_id'].split('-')[2]
    print(dictionary)

    # 创建运行程序所需要的线程池
    print("正在创建一个拥有5个线程的线程池")
    pool = tp.ThreadPool(5)
    template_change_files = []
    args_list_template = []
    if page2:
        template_change_files.extend(get_all_files(os.path.join('templates', dictionary['template_id'], '01管理手册')))
        template_change_files.extend(get_all_files(os.path.join('templates', dictionary['template_id'], '02程序文件')))
    else:
        template_change_files.extend(get_all_files(os.path.join('templates', dictionary['template_id'])))
    root_dir = os.path.join('templates', info_dict['template_id'])
    for file in template_change_files:
        args_list_template.append(([file, store_dir, root_dir, res_home, dictionary, info_dict], None))
    requests = tp.makeRequests(file_deal_process, args_list_template)
    for request in requests:
        pool.putRequest(request)
    pool.wait()
    check_file_list = []
    args_list_check = []

    # 复审模式相关逻辑的处理
    if re:
        tmppath = os.path.dirname(info_dict['tmp'])
        root_dir = tmppath
        check_file_list.extend(get_all_files(tmppath))
        for file in check_file_list:
            args_list_check.append(([file, store_dir, root_dir, res_home, old2new, info_dict], None))
        requests = tp.makeRequests(file_deal_process, args_list_check)
        for request in requests:
            pool.putRequest(request)
        pool.wait()

    exam_all(os.path.join(store_dir, res_home), 0)

    print('process done!')
